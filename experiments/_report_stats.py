"""Quick page/word stats for REPORT.docx."""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
doc = Document(str(ROOT / "results" / "REPORT.docx"))

total_paragraph_words = 0
para_count = 0
section_words: dict[str, int] = {}
current_section = "Pre-body"

for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt:
        continue
    para_count += 1
    n = len(txt.split())
    total_paragraph_words += n
    # Heading mu? (basit detection: bold first run + numara)
    if p.runs and p.runs[0].bold and (txt.startswith(("1.","2.","3.","4.","5.","6.","Referanslar","Grup","Özet"))):
        current_section = txt
        section_words[current_section] = 0
    else:
        section_words.setdefault(current_section, 0)
        section_words[current_section] += n

table_cells = 0
table_words = 0
table_words_per_table = []
for t in doc.tables:
    twc = 0
    for row in t.rows:
        for cell in row.cells:
            ct = cell.text.strip()
            if ct:
                table_cells += 1
                w = len(ct.split())
                table_words += w
                twc += w
    table_words_per_table.append(twc)

print(f"Paragraf sayisi      : {para_count}")
print(f"Paragraf kelime      : {total_paragraph_words}")
print(f"Tablo sayisi         : {len(doc.tables)}")
print(f"Tablo hucresi        : {table_cells}")
print(f"Tablo kelime         : {table_words}")
print(f"TOPLAM kelime        : {total_paragraph_words + table_words}")
print()
print("Bolum bazinda kelime:")
for k, v in section_words.items():
    print(f"  {k[:50]:50s}  {v} kelime")
print()
print("Tablo kelime/tablo:")
for i, w in enumerate(table_words_per_table):
    print(f"  Tablo {i}: {w} kelime")

# Inline image count
img_count = 0
for s in doc.inline_shapes:
    img_count += 1
print(f"\nInline gorsel sayisi : {img_count}")

# Sutun yapisi
print("\nSutun yapisi:")
ns_w = "w"
for i, s in enumerate(doc.sections):
    cols = s._sectPr.find(qn(f"{ns_w}:cols"))
    n = cols.get(qn(f"{ns_w}:num")) if cols is not None else "-"
    print(f"  bolum {i}: sutun sayisi = {n}")
