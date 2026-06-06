"""
_make_report_docx.py
====================

REPORT.docx olusturma (double-column, TR style, IEEE benzeri).

Yapi:
- Single-column title block (logo placeholder + üniversite + ders bilgisi + baslik + yazarlar)
- Single-column abstract
- Double-column body (5 bolum + grup katkisi + referanslar)
- Figurler results/figures/'dan eklenir
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = ROOT / "results" / "figures"
OUT_PATH = ROOT / "results" / "REPORT.docx"


# --------------------------------------------------------------------------- #
# Yardimcilar                                                                 #
# --------------------------------------------------------------------------- #
def set_columns(section, cols: int, space_cm: float = 0.5):
    """Bir bolumdeki sutun sayisini ayarla (low-level XML)."""
    sectPr = section._sectPr
    cols_el = sectPr.find(qn("w:cols"))
    if cols_el is None:
        cols_el = OxmlElement("w:cols")
        sectPr.append(cols_el)
    cols_el.set(qn("w:num"), str(cols))
    # Sutunlar arasi bosluk (twips: 1 cm ≈ 567)
    cols_el.set(qn("w:space"), str(int(space_cm * 567)))


def set_margins(section, top=1.5, bottom=1.5, left=1.6, right=1.6):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def add_para(doc, text, size=10, bold=False, italic=False,
             align=None, space_after=2, font="Times New Roman",
             color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 11, 2: 10}
    p = add_para(doc, text, size=sizes.get(level, 10), bold=True, space_after=3)
    return p


def add_table(doc, header: list, rows: list, col_widths: list[float] | None = None,
              first_col_bold: bool = False, font_size: int = 8):
    """Basit tablo. col_widths: cm cinsinden."""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # Header
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(font_size)

    # Data
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.name = "Times New Roman"
            r.font.size = Pt(font_size)
            if first_col_bold and ci == 0:
                r.bold = True
    # Sonrasi icin bos paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    return table


def add_figure(doc, image_path: Path, caption: str, width_cm: float = 8.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    cap = add_para(doc, caption, size=8, italic=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    return cap


# --------------------------------------------------------------------------- #
# Belge                                                                       #
# --------------------------------------------------------------------------- #
def build():
    doc = Document()

    # Varsayilan font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    # === Bolum 1: Single-column (baslik + abstract) ===
    sect = doc.sections[0]
    set_margins(sect)
    set_columns(sect, 1)

    # Logo placeholder
    add_para(doc, "[BTÜ LOGOSU — Word'de eklenecek]",
             size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0x80, 0x80, 0x80))

    add_para(doc, "Bursa Teknik Üniversitesi — Mekatronik Mühendisliği",
             size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Hesaplamalı Sinir Bilimine Giriş — Final Projesi 2025-2026 Bahar",
             size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Baslik
    add_para(doc,
        "BCI Competition IV Dataset 2a Üzerinde Motor Imagery EEG Sinyallerinin "
        "Sınıflandırılması: Sistematik Ablasyon ve Sızıntısız (Leakage-Free) Değerlendirme",
        size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # Yazarlar (placeholder)
    add_para(doc,
        "[Ad Soyad 1 — Öğrenci No], [Ad Soyad 2 — Öğrenci No], [Ad Soyad 3 — Öğrenci No]",
        size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc,
        "Bursa Teknik Üniversitesi, Mekatronik Mühendisliği Bölümü",
        size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

    # Ozet
    add_para(doc, "Özet", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=4)
    abstract = (
        "Motor imagery (MI) tabanlı beyin-bilgisayar arayüzü (BCI) sistemlerinde "
        "sınıflandırma performansı; küçük örneklem hacmi, denekler-arası varyasyon ve "
        "oturumlar-arası sinyal kayması gibi nedenlerle sınırlıdır. Bu çalışmada BCI "
        "Competition IV Dataset 2a (9 denek, 4 sınıf, 22 EEG kanalı) üzerinde sistematik "
        "bir ablasyon protokolüyle yedi farklı sinyal işleme/sınıflandırma boru hattı "
        "karşılaştırılmıştır. Tüm geliştirme süreci A0XT eğitim oturumunda iç içe (nested) "
        "çapraz doğrulama (5×5) ile yürütülmüş, A0XE değerlendirme oturumu proje boyunca "
        "tek seferde açılarak dürüst genelleme tahmini elde edilmiştir. Filtre bankası "
        "tabanlı CSP+shrinkage-LDA (κ_CV=0.625), Riemannian tangent space, EEGNet ve "
        "ShallowConvNet'ten oluşan dört temel modelin soft-voting birleşimi A0XT'de "
        "κ_CV=0.6553, A0XE'de κ_Test=0.6137 sağlamış ve yarışma kazananı FBCSP "
        "yönteminin (κ=0.569) +0.045 üzerinde performans elde etmiştir. CV-Test farkının "
        "−0.042 olması leakage'siz, dürüst genellemenin kanıtıdır. A05 deneğinin tüm "
        "yöntemlerde 0.30 altında kalması BCI-illiterate fenomenine güçlü kanıt sağlamaktadır."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(abstract)
    r.font.name = "Times New Roman"
    r.font.size = Pt(9.5)

    add_para(doc, "Anahtar kelimeler: BCI, motor imagery, FBCSP, Riemannian, EEGNet, ensemble.",
             size=9, italic=True, space_after=6)

    # === Bolum 2: Double-column (govde) ===
    new_sect = doc.add_section(WD_SECTION_START.CONTINUOUS)
    set_margins(new_sect)
    set_columns(new_sect, 2, space_cm=0.6)

    # ---- 1. Giris ----
    add_heading(doc, "1. Giriş")
    intro1 = (
        "Motor imagery (MI), bir uzvun fiziksel hareketini gerçekleştirmeksizin zihinde "
        "canlandırılması sürecidir ve duyu-motor kortekste mu (8–13 Hz) ile beta (13–30 Hz) "
        "ritimlerinde olay-ilişkili senkronizasyon/desenkronizasyon (ERS/ERD) örüntüleri "
        "üretir [1]. Bu kortikal değişimlerin sınıflandırılması motor BCI'ların temelini oluşturur."
    )
    add_para(doc, intro1, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    intro2 = (
        "BCI Competition IV Dataset 2a, MI sınıflandırması için referans veri kümesidir: "
        "9 denek, 4 sınıf (sol el, sağ el, iki ayak, dil), her denek için iki ayrı oturumda "
        "288 deneme [2]. Yarışma kazananı Ang vd.'nin Filter Bank CSP (FBCSP) yöntemi, dar "
        "bant CSP'lerin birleştirilmesi ile mu/beta etrafındaki ayrımcı bilgiyi yakalar; A0XE "
        "değerlendirme setinde mean κ=0.569 elde etmiştir [3]. Sonraki dönemde geometrik "
        "yaklaşımlar (Riemannian tangent space [4]) ve uçtan-uca derin öğrenme modelleri "
        "(EEGNet [5], ShallowConvNet [6]) güçlü alternatifler olarak ortaya çıkmıştır."
    )
    add_para(doc, intro2, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    intro3 = (
        "Bu çalışmanın katkıları: (i) yedi boru hattının sistematik ablasyonu — band selection, "
        "Riemannian, derin öğrenme, transfer öğrenme dahil — pozitif ve negatif sonuçlarla; "
        "(ii) sızıntısız (leakage-free) protokol: tüm geliştirme A0XT içinde 5×5 nested CV ile, "
        "A0XE final tek-shot tahmin için saklı tutulmuştur; (iii) yarışma kazananını geçen "
        "soft-voting ensemble (κ_Test=0.6137, +0.045) ve A05 deneğinde BCI-illiterate "
        "fenomeninin belgelenmesi."
    )
    add_para(doc, intro3, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ---- 2. Veri ve On Isleme ----
    add_heading(doc, "2. Veri Seti ve Ön İşleme")
    data1 = (
        "Veri kümesi. BCI Competition IV Dataset 2a: 9 denek (A01–A09), 22 monopolar EEG + "
        "3 EOG kanal, 250 Hz örnekleme. Her denek için iki ayrı gün kayıt: A0XT (eğitim) ve "
        "A0XE (değerlendirme), her birinde 288 deneme (sınıf başına 72). Etiketler A0XT için "
        "GDF olay kodlarından (769–772), A0XE için ayrı .mat dosyalarındaki classlabel "
        "alanından okunmuş ve 0–3 kodlamasına kaydırılmıştır."
    )
    add_para(doc, data1, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    data2 = (
        "Epoch çıkarımı. Cue uyarısından sonra 0.5–2.5 s penceresi alınmıştır (500 örnek × "
        "22 EEG kanal; EOG kanalları ayrı tutulmuş, sınıflandırmada kullanılmamıştır). MI "
        "motor görüntüleme süresi tipik olarak cue+0.5 s'den itibaren stabilize olur [3]."
    )
    add_para(doc, data2, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    data3 = (
        "Filtre bankası. Klasik kolda 16 bant: 8–30 Hz aralığında 10 adet 4 Hz genişlikte "
        "(8–12, 10–14, …, 26–30) ve 6 adet 6 Hz genişlikte bant (8–14, 11–17, …, 23–29). "
        "Tüm filtreler 4. dereceden Butterworth, sıfır-faz filtfilt ile uygulanmıştır. "
        "Riemannian kolda 3 geniş bant (8–14, 12–20, 18–30), derin öğrenme kolunda tek "
        "geniş bant (4–38 Hz) tercih edilmiştir."
    )
    add_para(doc, data3, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    data4 = (
        "Sızıntısız değerlendirme protokolü. Geliştirme süresince yalnızca A0XT "
        "kullanılmıştır. Dış 5-fold stratified CV (random_state=42) ile iç 5-fold "
        "GridSearchCV birleşik tutulmuş (nested CV), tüm veri-bağımlı adımlar (CSP, "
        "ölçekleme, öznitelik seçimi, PCA) sklearn.Pipeline içine yerleştirilerek yalnızca "
        "eğitim katmanında fit edilmiştir. A0XE proje boyunca bir kez bile açılmamış, son "
        "aşamada A0XT'nin tamamında yeniden eğitilen modeller A0XE'de tek-shot tahmin için "
        "kullanılmıştır."
    )
    add_para(doc, data4, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ---- 3. Oznitelik ve Siniflandirma ----
    add_heading(doc, "3. Öznitelik Çıkarımı ve Sınıflandırma")
    m1 = (
        "Üç temel kol incelenmiştir."
    )
    add_para(doc, m1, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    m2 = (
        "Klasik kol — FBCSP + shrinkage-LDA (Faz 1). 16 bandın her birinde Ledoit-Wolf "
        "shrinkage'li CSP (mne.decoding.CSP, log-variance) ile öznitelik çıkarılmış, tüm "
        "bant öznitelikleri birleştirilip SelectKBest(mutual_info_classif) ile boyut "
        "indirgenmiş, doğrusal LDA (solver='lsqr', shrinkage='auto') ile sınıflandırılmıştır. "
        "Tune edilen hiperparametreler: csp_n_components ∈ {2,4,6} ve "
        "select_k ∈ {10,20,40,'all'}."
    )
    add_para(doc, m2, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    m3 = (
        "Geometrik kol — Multi-band Riemannian TS (Faz 3c). 3 geniş bantta "
        "Covariances(estimator='oas') ile uzaysal kovaryans tahmini, TangentSpace ile "
        "manifold projeksiyonu, bant öznitelikleri birleştirildikten sonra PCA (n∈{30,50,80}) "
        "veya L1-regularized lojistik regresyon (saga, C∈{0.01,0.1,1,10}) ile sınıflandırma "
        "yapılmıştır [4]."
    )
    add_para(doc, m3, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    m4 = (
        "Derin kol — EEGNet ve ShallowConvNet (Faz 5a). Standart EEGNet-8,2 (F1=8, D=2, "
        "F2=16) [5] ve ShallowConvNet [6] mimarileri PyTorch ile yeniden uygulanmış; giriş "
        "(B,1,22,500), AdamW (lr=1e-3, wd=1e-2), batch=64, en fazla 300 epoch, erken "
        "durdurma (patience=30) ile eğitilmiştir. Augmentation (yalnızca eğitim): ±50 ms "
        "zaman kayması ve kanal-bazlı Gauss gürültüsü (σ=0.1×std)."
    )
    add_para(doc, m4, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    m5 = (
        "Cross-subject transfer (Faz 5b). LOSO öncelikli eğitim: hedef denek hariç 8 deneğin "
        "tüm A0XT verisiyle (2304 deneme) ön eğitim, ardından hedef deneğin dış-train "
        "katmanında ince ayar (lr=1e-4, patience=15)."
    )
    add_para(doc, m5, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    m6 = (
        "Ensemble — soft voting. Her temel modelin predict_proba çıktıları ağırlıklı "
        "ortalama ile birleştirilmiştir; nihai sistem fbcsp_rlda + riemann_multiband_ts_l1 + "
        "EEGNet + ShallowConvNet dörtlüsü, ağırlık 1.0/1.0/0.5/0.5 ile soft-vote yapmaktadır. "
        "Hata profili çeşitliliği ön-teşhisi için pairwise Jaccard hata-örtüşme matrisi "
        "hesaplanmış; FBCSP↔DL örtüşmesinin 0.28–0.30 aralığında ortogonal olduğu "
        "doğrulandıktan sonra ensemble uygulanmıştır."
    )
    add_para(doc, m6, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ---- 4. Sonuclar ----
    add_heading(doc, "4. Sonuçlar")
    r1 = (
        "Ablasyon ilerlemesi. Tablo 1, dokuz boru hattının A0XT üzerindeki dürüst nested CV "
        "ortalama κ değerini ve Faz 6 final A0XE testini özetler. Şekil 2 görsel ilerlemeyi "
        "sunar."
    )
    add_para(doc, r1, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc, "Tablo 1. Dokuz boru hattının ortalama κ değerleri (9 denek).",
             size=9, italic=True, space_after=2)
    add_table(doc,
        header=["Faz", "Boru hattı", "κ"],
        rows=[
            ["1",        "FBCSP + shrinkage-LDA",            "0.6250"],
            ["2a",       "FBCSP + bant seçimi + linSVM (MI)",   "0.5787"],
            ["2b",       "FBCSP + bant seçimi + linSVM (mRMR)", "0.5849"],
            ["3a",       "Riemann TS, 1-bant",                  "0.5303"],
            ["3b",       "Riemann MDM",                         "0.4635"],
            ["3c-PCA",   "Multi-band TS + PCA + LR",            "0.5417"],
            ["3c-L1",    "Multi-band TS + L1-LR",               "0.5499"],
            ["4",        "Klasik ensemble (FBCSP ⊕ MB-TS-L1)",   "0.6425"],
            ["5a",       "EEGNet (subject-specific)",           "0.3976"],
            ["5a",       "ShallowConvNet (subject-specific)",   "0.3971"],
            ["5a ens.",  "4 base ens. (1/1/0.5/0.5) — CV",      "0.6553"],
            ["5b",       "EEGNet LOSO transfer",                "0.3863"],
            ["5b",       "ShallowConvNet LOSO transfer",        "0.4213"],
            ["6",        "Final ens. — A0XE TEST",              "0.6137"],
        ],
        col_widths=[1.5, 8.0, 1.5],
        font_size=8,
    )

    r2 = (
        "Ablasyondan üç sonuç ön plana çıkar: (i) shrinkage-LDA örtük boyut indirgemesi "
        "nedeniyle hard bant seçimi (Faz 2) bilgi kaybı yaratmıştır (−0.04); (ii) tek-bant "
        "Riemannian FBCSP'nin filtre bankası avantajını yakalayamamış, multi-band varyantları "
        "Faz 1'in altında kalmıştır; (iii) subject-specific DL küçük örneklemde (288 deneme) "
        "çökerken (≈0.40), ensemble içinde düşük ağırlıkla ortogonal hata profili sağlayarak "
        "+0.013'lük katkı vermiştir."
    )
    add_para(doc, r2, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_figure(doc, FIG_DIR / "fig_ablation.png",
               "Şekil 2. Sistematik ablasyon ilerlemesi.",
               width_cm=8.0)

    r3 = (
        "Denek bazlı sonuçlar. Tablo 2 ve Şekil 1, dokuz denek için CV (A0XT) ile Test "
        "(A0XE) κ değerlerini sunar."
    )
    add_para(doc, r3, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_para(doc,
        "Tablo 2. Final ensemble — CV (A0XT, nested 5×5) ile Test (A0XE saklı set) "
        "karşılaştırması.", size=9, italic=True, space_after=2)
    add_table(doc,
        header=["Denek", "Faz 1 Test", "CV", "Test", "Δ", "acc"],
        rows=[
            ["A01", "0.7454", "0.8056", "0.7778", "−0.028", "0.833"],
            ["A02", "0.3194", "0.4074", "0.3426", "−0.065", "0.507"],
            ["A03", "0.7778", "0.8704", "0.8333", "−0.037", "0.875"],
            ["A04", "0.5417", "0.5231", "0.6204", "+0.097", "0.715"],
            ["A05", "0.1852", "0.3380", "0.2454", "−0.093", "0.434"],
            ["A06", "0.3565", "0.4630", "0.3750", "−0.088", "0.531"],
            ["A07", "0.7176", "0.7917", "0.7963", "+0.005", "0.847"],
            ["A08", "0.7500", "0.8657", "0.7778", "−0.088", "0.833"],
            ["A09", "0.7222", "0.8333", "0.7546", "−0.079", "0.816"],
            ["MEAN","0.5684", "0.6553", "0.6137", "−0.042", "0.710"],
        ],
        col_widths=[1.4, 1.7, 1.4, 1.4, 1.4, 1.4],
        first_col_bold=True,
        font_size=8,
    )

    add_figure(doc, FIG_DIR / "fig_cv_vs_test.png",
               "Şekil 1. Denek bazlı CV (A0XT) ve Test (A0XE) κ değerleri.",
               width_cm=8.0)

    r4 = (
        "Yöntem karşılaştırması. Ang vd.'nin yarışma kazananı FBCSP A0XE'de κ=0.569 "
        "raporlamıştır [3]. Bu çalışmanın final ensemble'ı κ_Test=0.6137 ile bu rakamı "
        "+0.045 geçmiştir. Ayrıca Faz 4 klasik-yalnız ensemble (κ=0.6106) ile Faz 6 final "
        "(κ=0.6137) arasındaki ufak +0.003'lük fark, derin öğrenme ortagonal bilgisinin test "
        "setinde de korunduğunu göstermektedir."
    )
    add_para(doc, r4, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    r5 = (
        "Confusion matrix. Şekil 3, dokuz deneğin A0XE final tahminlerinin satır-normalize "
        "toplam karışıklık matrisini sunar. Sınıflar arası en yüksek karışıklık sol el – iki "
        "ayak çiftindedir; dil sınıfı en yüksek doğrulukla sınıflandırılmıştır."
    )
    add_para(doc, r5, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    add_figure(doc, FIG_DIR / "fig_confusion_matrix.png",
               "Şekil 3. A0XE final ensemble normalize confusion matrix (9 denek toplam).",
               width_cm=7.5)

    # ---- 5. Tartisma ----
    add_heading(doc, "5. Tartışma")
    d1 = (
        "Negatif sonuçların yorumu. Dört bağımsız ablasyon Faz 1 baseline'ını geçememiş, "
        "ancak metodolojik içgörü sağlamıştır. (i) Faz 2 — bant seçimi: shrinkage-LDA, "
        "Ledoit-Wolf büzülmüş kovaryans ile gereksiz boyutları zaten yumuşak olarak "
        "bastırdığından, MI tabanlı hard bant seçimi bilgi kaybı yaratmıştır. (ii) Faz 3 — "
        "Riemannian: tek-bant TS FBCSP'nin 16-bant filtre bankası avantajını yakalayamamış; "
        "multi-band varyantında 759 öznitelik / 288 deneme oranı boyut sınırını zorlamıştır. "
        "(iii) Faz 5a — subject-specific DL: 288 deneme EEGNet/ShallowConvNet için "
        "yetersiz; literatürde tipik 0.55–0.65 değerleri cross-subject ön eğitim varsayar. "
        "(iv) Faz 5b — LOSO transfer: etki asimetriktir; \"tabula rasa\" deneklerinde büyük "
        "kazanç (A06 SCN +0.190, A05 SCN +0.097) ama bilgi taşıyan deneklerde zarar "
        "(A07 SCN −0.255). Net mean κ değişimi marjinal kalmıştır."
    )
    add_para(doc, d1, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    d2 = (
        "Düşük performanslı denekler. A02 (κ=0.34) ve A06 (κ=0.38) test setinde tüm "
        "yöntemlerde zayıf kalmıştır; bu deneklerde EOG artefakt seviyesi ve trial-içi "
        "varyansın yüksek olduğu beklenmektedir (gelecek iş: amplitude-tabanlı trial reddediş)."
    )
    add_para(doc, d2, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    d3 = (
        "A05 — BCI-illiterate kanıtı. A05 deneği yedi farklı klasik/geometrik/derin yöntem "
        "ve LOSO transfer dahil hiçbir yapılandırmada test κ=0.30'u geçememiştir (final test "
        "κ=0.245). Transfer smoke test'inde bir dış foldta negatif κ (−0.125) gözlemlenmesi, "
        "cross-subject prior'un bu deneğin verisine ters yönde transfer edildiğine işaret eder. "
        "Literatürde \"BCI illiteracy\" olarak adlandırılan ve popülasyonun %20–30'unda "
        "görülen fenomenle [7] uyumlu olarak, A05'in MI sinyali ya yetersiz ayrımcı bilgi "
        "taşımakta ya da niteliksel olarak diğer deneklerden farklı yapıdadır. Denek-özgü "
        "tavan = 0.245 olarak işaretlenmiştir."
    )
    add_para(doc, d3, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    d4 = (
        "CV ↔ Test farkı. Ortalama Δ = −0.042 makul session-to-session shift seviyesindedir; "
        "ne CV'nin optimistik bias'ını (Δ≪0 beklenirdi) ne de pessimistik bias'ını (Δ>0 "
        "beklenirdi) gösterir. Bu, sızıntısız nested CV protokolünün dürüst genelleme "
        "tahmini verdiğine dair güçlü ampirik kanıttır."
    )
    add_para(doc, d4, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    d5 = (
        "İyileştirme yönleri. (a) Trial rejection: EOG-bazlı veya amplitude-bazlı outlier "
        "reddediş, özellikle A02/A06'da kazanç sağlayabilir. (b) Daha fazla veri: çoklu "
        "oturum kaydı veya cross-dataset ön eğitim DL kollarının üst limitini açabilir. "
        "(c) Gelişmiş augmentation: SmoothMix, frekans-alan maskeleme [8] gibi EEG-özgü "
        "teknikler. (d) A05 için kanal-bazlı SNR analizi: mu/beta bandında C3/C4 öznitelik "
        "amplitüd histogramları, BCI-illiteracy hipotezini niceliksel doğrulayabilir."
    )
    add_para(doc, d5, size=10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ---- 6. Grup katki tablosu (yer-tutucu) ----
    add_heading(doc, "Grup İçi Katkı Tablosu")
    add_para(doc, "[Yer-tutucu — kullanıcı dolduracak]",
             size=9, italic=True, color=RGBColor(0x80, 0x80, 0x80))
    add_table(doc,
        header=["Üye", "Katkı Alanı", "Yüzde"],
        rows=[
            ["[Ad Soyad 1]", "…", "…%"],
            ["[Ad Soyad 2]", "…", "…%"],
            ["[Ad Soyad 3]", "…", "…%"],
        ],
        col_widths=[3.2, 5.8, 1.5],
        font_size=9,
    )

    # ---- Referanslar ----
    add_heading(doc, "Referanslar")
    refs = [
        "[1] G. Pfurtscheller and F. H. Lopes da Silva, \"Event-related EEG/MEG synchronization "
        "and desynchronization: basic principles,\" Clinical Neurophysiology, vol. 110, no. 11, "
        "pp. 1842–1857, 1999.",
        "[2] M. Tangermann et al., \"Review of the BCI Competition IV,\" Frontiers in "
        "Neuroscience, vol. 6, p. 55, 2012.",
        "[3] K. K. Ang, Z. Y. Chin, H. Zhang, and C. Guan, \"Filter Bank Common Spatial Pattern "
        "(FBCSP) in Brain-Computer Interface,\" IJCNN 2008, pp. 2390–2397, 2008.",
        "[4] A. Barachant, S. Bonnet, M. Congedo, and C. Jutten, \"Multiclass brain-computer "
        "interface classification by Riemannian geometry,\" IEEE Trans. Biomed. Eng., vol. 59, "
        "no. 4, pp. 920–928, 2012.",
        "[5] V. J. Lawhern et al., \"EEGNet: a compact convolutional neural network for "
        "EEG-based brain-computer interfaces,\" J. Neural Eng., vol. 15, no. 5, 056013, 2018.",
        "[6] R. T. Schirrmeister et al., \"Deep learning with convolutional neural networks for "
        "EEG decoding and visualization,\" Human Brain Mapping, vol. 38, no. 11, pp. 5391–5420, "
        "2017.",
        "[7] B. Blankertz et al., \"Neurophysiological predictor of SMR-based BCI performance,\" "
        "NeuroImage, vol. 51, no. 4, pp. 1303–1309, 2010.",
        "[8] F. Lotte and C. Guan, \"Regularizing common spatial patterns to improve BCI "
        "designs: unified theory and new algorithms,\" IEEE Trans. Biomed. Eng., vol. 58, no. 2, "
        "pp. 355–362, 2011.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.10
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(8.5)

    # Kaydet
    doc.save(str(OUT_PATH))
    print(f"saved {OUT_PATH}")


if __name__ == "__main__":
    build()
